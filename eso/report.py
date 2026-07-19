"""Optional Word-document report showing original vs. optimized sequences
with per-nucleotide differences highlighted.
"""

from os import path

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def create_word_document_with_highlighted_differences(sequences_data, output_path):
    """
    Parameters
    ----------
    sequences_data: list of (sequence_name, original_seq, final_seq) tuples.
    output_path: directory in which to save 'sequence_comparison.docx'.
    """
    if not DOCX_AVAILABLE:
        print("python-docx not available (install the 'docx-report' extra), skipping Word document generation")
        return

    doc = Document()
    doc.add_heading('Sequence Optimization Results', 0)

    for seq_name, original_seq, final_seq in sequences_data:
        doc.add_heading(f'Sequence: {seq_name}', level=1)

        doc.add_heading('Original Sequence:', level=2)
        original_paragraph = doc.add_paragraph()

        doc.add_heading('Final Sequence:', level=2)
        final_paragraph = doc.add_paragraph()

        for i, char in enumerate(original_seq):
            run = original_paragraph.add_run(char)
            if i < len(final_seq) and char != final_seq[i]:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        for i, char in enumerate(final_seq):
            run = final_paragraph.add_run(char)
            if i < len(original_seq) and char != original_seq[i]:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if seq_name != sequences_data[-1][0]:
            doc.add_page_break()

    doc_path = path.join(output_path, 'sequence_comparison.docx')
    doc.save(doc_path)
    print(f"Word document saved to: {doc_path}")
