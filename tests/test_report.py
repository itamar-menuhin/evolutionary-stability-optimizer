"""Tests for eso.report.create_word_document_with_highlighted_differences -
previously untested.

Includes a regression test for a real bug found while reviewing: the
page-break-between-entries logic compared `seq_name != sequences_data[-1][0]`
(by name) instead of by position, so an entry whose name happened to match
the true last entry's name (e.g. two files/records that coincidentally share
a stem) would wrongly skip its own page break. Confirmed directly: with
three entries where the first and third share a name, only 1 page break was
produced instead of the expected 2.
"""

import pytest

docx = pytest.importorskip("docx", reason="python-docx not installed (optional 'docx-report' extra)")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from eso.report import create_word_document_with_highlighted_differences  # noqa: E402


def _count_page_breaks(doc):
    breaks = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    breaks += 1
    return breaks


def _highlighted_chars(paragraph):
    return [run.text for run in paragraph.runs if run.font.highlight_color is not None]


def test_creates_the_expected_file(tmp_path):
    create_word_document_with_highlighted_differences(
        [('gene_a', 'ACGT', 'ACGA')], str(tmp_path))
    assert (tmp_path / 'sequence_comparison.docx').exists()


def test_page_break_after_every_entry_except_the_last(tmp_path):
    data = [
        ('a', 'AAAA', 'AAAT'),
        ('b', 'CCCC', 'CCCG'),
        ('c', 'GGGG', 'GGGT'),
    ]
    create_word_document_with_highlighted_differences(data, str(tmp_path))
    doc = Document(str(tmp_path / 'sequence_comparison.docx'))
    assert _count_page_breaks(doc) == 2


def test_page_break_logic_is_not_confused_by_a_repeated_name(tmp_path):
    # regression test: the first and third entries share a name ('gene'),
    # which used to make the first entry wrongly skip its page break.
    data = [
        ('gene', 'AAAA', 'AAAT'),
        ('other', 'CCCC', 'CCCG'),
        ('gene', 'GGGG', 'GGGT'),
    ]
    create_word_document_with_highlighted_differences(data, str(tmp_path))
    doc = Document(str(tmp_path / 'sequence_comparison.docx'))
    assert _count_page_breaks(doc) == 2


def test_single_entry_has_no_page_break(tmp_path):
    create_word_document_with_highlighted_differences(
        [('only_one', 'AAAA', 'AAAT')], str(tmp_path))
    doc = Document(str(tmp_path / 'sequence_comparison.docx'))
    assert _count_page_breaks(doc) == 0


def test_differing_positions_are_highlighted_in_both_sequences(tmp_path):
    create_word_document_with_highlighted_differences(
        [('gene', 'ACGTACGT', 'ACGAACCT')], str(tmp_path))
    doc = Document(str(tmp_path / 'sequence_comparison.docx'))

    # the two sequence-body paragraphs use the "Normal" style; headings/title don't.
    body_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Normal']
    original_paragraph, final_paragraph = body_paragraphs[0], body_paragraphs[1]

    assert _highlighted_chars(original_paragraph) == ['T', 'G']
    assert _highlighted_chars(final_paragraph) == ['A', 'C']


def test_identical_sequences_have_no_highlights(tmp_path):
    create_word_document_with_highlighted_differences(
        [('gene', 'ACGT', 'ACGT')], str(tmp_path))
    doc = Document(str(tmp_path / 'sequence_comparison.docx'))

    body_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Normal']
    for paragraph in body_paragraphs:
        assert _highlighted_chars(paragraph) == []
